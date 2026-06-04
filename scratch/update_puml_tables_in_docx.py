import docx
from docx.shared import Pt, RGBColor
import os

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
puml_dir = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/diagrams"

doc = docx.Document(doc_path)

mapping = {
    "UC_Overall": "2.01_uc_overall.puml",
    "UC_Auth": "2.02_uc_auth.puml",
    "UC_Room": "2.03_uc_room.puml",
    "UC_Contract": "2.04_uc_contract.puml",
    "UC_Billing": "2.05_uc_billing.puml",
    "UC_Report": "2.06_uc_report.puml",
    "UC_Utilities": "2.06b_uc_utilities.puml",
    "Act_Register_OTP": "2.07_act_register_otp.puml",
    "Act_Room": "2.08_act_room.puml",
    "Act_Contract_Sign": "2.09_act_contract_sign.puml",
    "Act_Meter_Invoice": "2.10_act_meter_invoice.puml",
    "Act_Payment": "2.11_act_payment.puml",
    "ClassDiagram": "2.12_class_diagram.puml",
    "Seq_Register_OTP": "2.13_seq_register_otp.puml",
    "Seq_Contract_Sign": "2.14_seq_contract_sign.puml",
    "Seq_Meter_Invoice": "2.15_seq_meter_invoice.puml",
    "Seq_Payment_VNPay": "2.16_seq_payment_vnpay.puml",
    "Seq_Chatbot_AI": "2.17_seq_chatbot_ai.puml",
    "Architecture_Deployment": "2.18_architecture.puml"
}

def style_cell_run(run, text):
    run.text = text
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(80, 80, 80)

print("Scanning tables in docx to update PUML codes...")
updated_count = 0

for idx, table in enumerate(doc.tables):
    if len(table.rows) == 1 and len(table.columns) == 1:
        cell_text = table.cell(0, 0).text
        if "@startuml" in cell_text:
            # Determine mapping key
            matched_key = None
            for key in mapping.keys():
                if f"@startuml {key}" in cell_text or f"@startuml\n{key}" in cell_text:
                    matched_key = key
                    break
            
            if not matched_key:
                # Try fallback line search
                for line in cell_text.splitlines():
                    if "@startuml" in line:
                        parts = line.split()
                        if len(parts) > 1:
                            potential_key = parts[1]
                            if potential_key in mapping:
                                matched_key = potential_key
                                break
            
            if matched_key:
                puml_file = mapping[matched_key]
                puml_path = os.path.join(puml_dir, puml_file)
                if os.path.exists(puml_path):
                    with open(puml_path, "r", encoding="utf-8") as f:
                        new_code = f.read().strip()
                    
                    cell = table.cell(0, 0)
                    cell.paragraphs[0].text = ""
                    style_cell_run(cell.paragraphs[0].add_run(), new_code)
                    print(f"✅ Table {idx:02d} updated with {puml_file}")
                    updated_count += 1
                else:
                    print(f"⚠️ PUML file not found: {puml_path}")
            else:
                print(f"❓ Could not map table {idx} containing: '{cell_text[:50]}...'")

if updated_count > 0:
    doc.save(doc_path)
    print(f"🎉 Success! Updated {updated_count} tables in the Word document.")
else:
    print("No tables were updated.")
