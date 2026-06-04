import docx

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
doc = docx.Document(doc_path)

print("=== Table 06 Row 3 Col 1 ===")
print(doc.tables[6].rows[3].cells[1].text)

print("\n=== Table 06 Row 6 Col 1 ===")
print(doc.tables[6].rows[6].cells[1].text)

print("\n=== Table 16 Row 8 ===")
row_cells = doc.tables[16].rows[8].cells
for c_idx, cell in enumerate(row_cells):
    print(f"Col {c_idx}: '{cell.text}'")
