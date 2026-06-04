import docx

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
doc = docx.Document(doc_path)

print("=== Scanning Paragraphs for SMS/Zalo/OTP ===")
for idx, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if any(term in txt.lower() for term in ["zalo", "sms", "otp"]):
        print(f"P {idx:03d} | '{txt[:150]}'")

print("\n=== Scanning Tables for SMS/Zalo/OTP ===")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell_txt = cell.text.strip()
            if any(term in cell_txt.lower() for term in ["zalo", "sms", "otp"]):
                first_line = cell_txt.split("\n")[0]
                print(f"Table {t_idx:02d} [Row {r_idx}, Col {c_idx}] | '{first_line[:120]}'")
                break
