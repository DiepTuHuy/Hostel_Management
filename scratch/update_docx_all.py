import docx
from docx.shared import Pt, RGBColor
import os

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
doc = docx.Document(doc_path)

puml_dir = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/diagrams"

def get_puml_code(filename):
    path = os.path.join(puml_dir, filename)
    with open(path, "r", encoding="utf-8") as f:
        return f.read().strip()

def style_cell_run(run, text):
    run.text = text
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(80, 80, 80)

# 1. Update PUML Code Blocks in Table 10 and Table 13
print("Updating PUML code block tables...")
contract_code = get_puml_code("2.04_uc_contract.puml")
cell_10 = doc.tables[10].cell(0, 0)
cell_10.paragraphs[0].text = ""
style_cell_run(cell_10.paragraphs[0].add_run(), contract_code)
print("✅ Updated Table 10 (2.04_uc_contract.puml)")

utilities_code = get_puml_code("2.06b_uc_utilities.puml")
cell_13 = doc.tables[13].cell(0, 0)
cell_13.paragraphs[0].text = ""
style_cell_run(cell_13.paragraphs[0].add_run(), utilities_code)
print("✅ Updated Table 13 (2.06b_uc_utilities.puml)")

# 2. Update Table 16: Delete Row 8
print("Updating Table 16...")
table_16 = doc.tables[16]
if len(table_16.rows) > 8:
    row_to_delete = table_16.rows[8]
    table_16._tbl.remove(row_to_delete._tr)
    print("✅ Deleted Row 8 in Table 16")

# 3. Update Text in other tables
print("Updating Table 02, Table 03, Table 06...")
# Table 2 Row 4 Col 1
t2_c = doc.tables[2].rows[4].cells[1]
t2_c.text = "Tạo hồ sơ khách thuê, ký hợp đồng (có ký số trực tuyến), gia hạn, sửa đổi, chấm dứt hợp đồng."
for p in t2_c.paragraphs:
    for run in p.runs:
        run.font.name = 'Times New Roman'

# Table 3 Row 8 Col 1
t3_c = doc.tables[3].rows[8].cells[1]
t3_c.text = "Lưu hoá đơn điện tử theo Nghị định 123/2020/NĐ-CP."
for p in t3_c.paragraphs:
    for run in p.runs:
        run.font.name = 'Times New Roman'

# Table 6 Row 3 Col 1
t6_r3 = doc.tables[6].rows[3].cells[1]
t6_r3.text = "UC15 Thêm hồ sơ khách; UC16 Lập hợp đồng thuê; UC17 Ký số / xác nhận hợp đồng (v1 giả lập – lưu URL PDF, v2 nhúng chữ ký CA); UC18 Gia hạn hợp đồng; UC19 Sửa đổi hợp đồng; UC20 Chấm dứt hợp đồng / trả phòng."
for p in t6_r3.paragraphs:
    for run in p.runs:
        run.font.name = 'Times New Roman'

# Table 6 Row 6 Col 1
t6_r6 = doc.tables[6].rows[6].cells[1]
t6_r6.text = "UC37 Gửi thông báo tự động qua email (v1 hỗ trợ kênh email Gmail SMTP với 4 hàm: OTP đăng ký, OTP quên mật khẩu, thông báo hợp đồng, nhắc nợ; kênh SMS/Zalo đẩy sang v2); UC38 Tìm kiếm phòng (khách vãng lai); UC39 Đặt cọc giữ phòng online; UC41 (mới) Trợ lý ảo AI Chatbot có chế độ Offline Fallback (Google Gemini 2.5 Flash + bộ luật cứng đọc trực tiếp MongoDB khi Gemini API gặp rate limit)."
for p in t6_r6.paragraphs:
    for run in p.runs:
        run.font.name = 'Times New Roman'

# 4. Update Paragraphs
print("Updating Paragraphs...")
# P 280
p_280 = doc.paragraphs[280]
p_280.text = 'Xác định rõ 04 tác nhân và 39 ca sử dụng (bao gồm UC41 Chatbot AI mới bổ sung) được tổ chức thành 6 nhóm chức năng chính, đảm bảo bao phủ toàn bộ nghiệp vụ vận hành. Đã triển khai backend thực tế đạt 35/39 UC (~90%); 3 UC ở mức giả lập có tích hợp dịch vụ thật ở v2 (UC17 ký số, UC27 VNPay, UC37 SMS/Zalo); 1 UC đẩy về Hướng phát triển ở Mục 4.3 (UC36 xuất Excel/PDF).'
for run in p_280.runs:
    run.font.name = 'Times New Roman'

# P 284
p_284 = doc.paragraphs[284]
p_284.text = 'Đề xuất phương án tích hợp cổng thanh toán (VNPay/MoMo/QR) và ký số trực tuyến - vốn là các "pain-point" lớn của quy trình thủ công.'
for run in p_284.runs:
    run.font.name = 'Times New Roman'

# P 290
p_290 = doc.paragraphs[290]
p_290.text = 'Tự động hoá tối đa các tác vụ lặp lại: tính tiền, sinh hoá đơn, gửi nhắc nợ.'
for run in p_290.runs:
    run.font.name = 'Times New Roman'

# P 311 (Delete UC40 first, to not shift index 308)
p_311 = doc.paragraphs[311]
p_311.text = ""
p_311._p.getparent().remove(p_311._p)
print("✅ Deleted P 311 (UC40)")

# P 308 (Delete UC21)
p_308 = doc.paragraphs[308]
p_308.text = ""
p_308._p.getparent().remove(p_308._p)
print("✅ Deleted P 308 (UC21)")

doc.save(doc_path)
print("🎉 Success! Word document text and tables fully updated!")
