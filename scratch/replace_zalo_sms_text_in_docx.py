import docx
import re

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
doc = docx.Document(doc_path)

replacements = [
    # OTP over SMS removals
    (", OTP qua SMS", ""),
    ("OTP qua SMS", "OTP qua Email"),
    ("OTP gửi qua email/SMS", "OTP gửi qua email"),
    ("OTP qua email/SMS", "OTP qua email"),
    ("mã OTP gửi qua email/SMS", "mã OTP gửi qua email"),
    ("xác thực 2 lớp (OTP qua SMS/email)", "xác thực 2 lớp (OTP qua email)"),
    ("OTP qua SMS/email", "OTP qua email"),
    
    # Zalo OA / SMS notifications -> Telegram Bot
    ("SMS/Zalo/Email", "Telegram/Email"),
    ("Email/Zalo/SMS", "Email/Telegram"),
    ("email/Zalo/SMS", "email/Telegram"),
    ("Zalo/SMS/Email", "Telegram/Email"),
    ("Zalo/SMS", "Telegram"),
    ("Zalo OA & SMS Brandname", "Telegram Bot"),
    ("Zalo OA \u0026 SMS Brandname", "Telegram Bot"),
    ("Zalo OA SDK và eSMS.vn Brandname để bổ sung 2 kênh SMS và Zalo", "Telegram Bot API"),
    ("Zalo OA", "Telegram Bot"),
    ("Zalo", "Telegram"),
    ("SMS Brand", "Telegram Bot"),
    ("SMS", "Telegram"),
    
    # Specific phrasing fixes
    ("kênh Telegram Bot đẩy sang v2", "kênh Telegram Bot"),
    ("kênh Telegram/Telegram đẩy sang v2", "kênh Telegram"),
]

def replace_text(text):
    for old, new in replacements:
        text = text.replace(old, new)
    return text

print("=== Replacing text in Paragraphs ===")
p_count = 0
for idx, p in enumerate(doc.paragraphs):
    original_text = p.text
    new_text = replace_text(original_text)
    if original_text != new_text:
        # Preserve styling as much as possible
        # Simple text replacement directly inside the paragraph
        # docx has runs, so to keep formatting we can replace text inside runs,
        # but string matching across runs is tricky.
        # Let's check if the paragraph contains simple text or multiple runs.
        # A simple way to replace text while preserving runs if they are split is to replace inside runs or replace entire paragraph text.
        # Since we want to preserve font name (Times New Roman) and font sizes, let's update text of runs
        # or clear and re-add.
        # Let's try to do it by reconstructing runs if necessary, or just replacing text in runs that contain the match.
        # Alternatively, we can assign p.text = new_text and set font name of runs.
        # Let's look if we can just do a robust run text replacement:
        combined = "".join(r.text for r in p.runs)
        if replace_text(combined) != combined:
            # Clear all runs and add a single run with new_text styled with Times New Roman
            p.text = ""
            run = p.add_run(new_text)
            run.font.name = 'Times New Roman'
            print(f"P {idx:03d} | '{original_text[:80]}...' -> '{new_text[:80]}...'")
            p_count += 1

print("\n=== Replacing text in Tables ===")
t_count = 0
for t_idx, table in enumerate(doc.tables):
    # Skip PUML code tables
    if len(table.rows) == 1 and len(table.columns) == 1 and "@startuml" in table.cell(0, 0).text:
        continue
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            # Check cell paragraphs
            cell_changed = False
            for p in cell.paragraphs:
                original_text = p.text
                new_text = replace_text(original_text)
                if original_text != new_text:
                    p.text = ""
                    run = p.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    cell_changed = True
            if cell_changed:
                print(f"Table {t_idx:02d} [Row {r_idx}, Col {c_idx}] updated")
                t_count += 1

doc.save(doc_path)
print(f"\n🎉 Success! Replaced terms in {p_count} paragraphs and {t_count} table cells in the Word document.")
