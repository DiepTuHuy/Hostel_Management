import docx

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
doc = docx.Document(doc_path)

indices = [280, 284, 290, 308, 311]
for idx in indices:
    if idx < len(doc.paragraphs):
        print(f"P {idx:03d}: '{doc.paragraphs[idx].text}'")
