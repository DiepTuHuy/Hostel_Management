import docx

doc_path = "/Users/dieptuhuy/Library/CloudStorage/GoogleDrive-dieptuhuy80@gmail.com/Other computers/My Computer 3/D:/Study/System_Design/docs/Báo cáo PTTKHT - Quản lý chuỗi nhà trọ (đã chỉnh sửa).docx"
doc = docx.Document(doc_path)

print(f"Total tables: {len(doc.tables)}")
for idx, table in enumerate(doc.tables):
    # Check if table has 1 cell and contains plantuml code
    if len(table.rows) == 1 and len(table.columns) == 1:
        text = table.cell(0, 0).text
        if "@startuml" in text:
            # Look at title or first few lines
            lines = text.split("\n")
            title_line = next((l for l in lines if "title" in l or "@startuml" in l), "")
            print(f"Table {idx:02d} | title: '{title_line.strip()}' | characters: {len(text)}")
